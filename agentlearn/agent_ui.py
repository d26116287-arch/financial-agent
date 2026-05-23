# 国内网络：用 HuggingFace 镜像
import os
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
import openpyxl
from langchain_openai import ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.tools import tool
from langchain.agents import create_agent
from langgraph.checkpoint.memory import InMemorySaver
from langchain_community.tools.tavily_search import TavilySearchResults
import easyocr
from docx import Document as DocxDocument
import io

# ============================================================
# 导出 Word 文档
# ============================================================

def build_docx(messages):
    """把对话历史转成 Word 文件，返回 bytes 供下载。"""
    doc = DocxDocument()
    doc.add_heading("财务分析对话记录", level=1)

    for msg in messages:
        role = msg["role"]
        content = msg["content"]
        label = "用户" if role == "user" else "分析助手"
        doc.add_heading(label, level=3)
        doc.add_paragraph(content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ============================================================
# 创建 Agent（普通函数，用 st.session_state 避免重复创建）
# 不再用 @st.cache_resource —— 那个会在页面加载时同步跑，卡住整个页面
# 改成懒加载：用户第一次发问题时才初始化
# ============================================================

def create_my_agent(document_path):
    """创建 Agent，耗时操作。调用一次后存入 session_state，不再重复执行。"""
    status = st.status("正在初始化...", expanded=True)

    status.update(label="正在加载 DeepSeek LLM...")
    llm = ChatOpenAI(
        api_key=st.secrets["DEEPSEEK_API_KEY"],
        base_url="https://api.deepseek.com",
        model="deepseek-v4-pro",
        max_tokens=4096,
        extra_body={"thinking": {"type": "disabled"}},
    )

    status.update(label="正在加载向量模型...")
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-base-zh-v1.5")

    status.update(label="正在读取本地文档并建立向量库...")
    if document_path.endswith(".pdf"):
        loader = PyPDFLoader(document_path)
        documents = loader.load()
    elif document_path.endswith((".xls", ".xlsx")):
        wb = openpyxl.load_workbook(document_path)
        documents = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines = []
            for row in ws.iter_rows(values_only=True):
                line = " | ".join(str(cell) if cell is not None else "" for cell in row)
                lines.append(line)
            text = "\n".join(lines)
            documents.append(Document(
                page_content=text,
                metadata={"source": document_path, "sheet": sheet_name}
            ))
    chunks = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100).split_documents(documents)
    vectorstore = Chroma.from_documents(chunks, embeddings, collection_name="agent_ui_collection")
    retriever = vectorstore.as_retriever(search_kwargs={"k": 6})

    @tool
    def search_local_docs(query: str) -> str:
        """搜索本地已上传的财务文档。"""
        docs = retriever.invoke(query)
        if not docs:
            return "未在本地文档中找到相关内容。"
        return "\n\n---\n\n".join(
            f"[来源: {doc.metadata.get('source', '未知')}]\n{doc.page_content}"
            for doc in docs
        )

    status.update(label="正在加载网络搜索工具...")
    _tavily = TavilySearchResults(
        tavily_api_key=st.secrets["TAVILY_API_KEY"],
        max_results=5,
        search_depth="advanced",
    )
    @tool
    def search_web(query: str) -> str:
        """搜索互联网获取最新公开信息。"""
        try:
            return _tavily.invoke(query)
        except Exception as e:
            return f"网络搜索失败（{type(e).__name__}）。"

    status.update(label="正在加载 OCR 模型...")
    _ocr = easyocr.Reader(["ch_sim", "en"], gpu=False)

    @tool
    def analyze_image(image_path: str) -> str:
        """识别图片中的文字（OCR）。"""
        try:
            results = _ocr.readtext(image_path)
            if not results:
                return "图片中未识别到文字。"
            lines = [f"  {text}  (置信度: {conf:.0%})" for (_, text, conf) in results]
            return "[从图片中识别的文字]:\n" + "\n".join(lines)
        except FileNotFoundError:
            return f"找不到图片文件: {image_path}"
        except Exception as e:
            return f"图片识别失败（{type(e).__name__}）: {str(e)}"

    status.update(label="正在组装 Agent...")
    SYSTEM_PROMPT = """你是一个严谨的财务分析助手，拥有三个工具：
    1. search_local_docs —— 搜索本地已上传的财务文档
    2. search_web —— 搜索互联网获取最新公开信息
    3. analyze_image —— 识别图片中的文字（OCR）
    选择工具的规则：
    - 公司内部数据、历史报销、合同条款 → search_local_docs
    - 最新政策、税法变动、行业趋势 → search_web
    - 用户提到图片文件（jpg/png）、扫描件、发票照片 → analyze_image
    - 不确定时多工具配合，综合回答
    - 严格基于工具结果回答，找不到就如实说
    - 用中文回答，清晰有条理"""

    agent = create_agent(
        llm,
        [search_local_docs, search_web, analyze_image],
        system_prompt=SYSTEM_PROMPT,
        checkpointer=InMemorySaver(),
    )

    status.update(label="初始化完成！", state="complete")
    return agent

# ============================================================
# 页面 UI
# ============================================================

st.title("📊 财务分析助手")

# -- 侧边栏 --
with st.sidebar:
    st.header("📁 文档管理")
    uploaded_file = st.file_uploader("上传文件", type=["pdf", "xls", "xlsx"])
    if uploaded_file is not None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        document_path = os.path.join(script_dir, uploaded_file.name)
        # 只有当文件真的换了，才重新处理（避免每次请求都重置）
        if st.session_state.get("Document") != document_path:
            with open(document_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.success(f"已上传并保存文件: {uploaded_file.name}")
            # 记住当前使用的文档路径，后续创建 Agent 时用它
            st.session_state.Document = document_path
            # 清空旧 Agent，下次发问题时用新文档重建
            st.session_state.agent = None
            # 清空聊天记录，因为文档变了，之前的对话上下文没意义
            st.session_state.messages = []

    st.divider()
    st.header("📤 导出")
    if st.session_state.get("messages"):
        st.download_button(
            label="导出为 Word 文档",
            data=build_docx(st.session_state.messages),
            file_name="财务分析记录.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# ---------- ① 初始化持久化数据 ----------
# session_state 是字典，跨请求存在。普通变量每次请求都重置。
if "Document" not in st.session_state:
    st.session_state.Document = None        # 文档对象（用户上传后存入）
if "agent" not in st.session_state:
    st.session_state.agent = None        # Agent 实例（懒加载，第一次发问才创建）

if "messages" not in st.session_state:
    st.session_state.messages = []       # 对话历史，格式：[{"role":"user","content":"..."}, ...]

# ---------- ② 渲染历史对话 ----------
# 每次请求重跑时，先把之前所有气泡画出来。第①次加载列表为空，啥也不画。
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

# ---------- ③ 接收新输入 ----------
question = st.chat_input("请输入你的问题：")

if question:
    # ---------- ④ 显示用户气泡 + 存入历史 ----------
    with st.chat_message("user"):
        st.write(question)
    st.session_state.messages.append({"role": "user", "content": question})

    # ---------- ⑤ 决定用哪个文档，然后按需创建 Agent ----------
    # 优先使用用户上传的文档路径，没有上传才用默认 example.pdf
    if st.session_state.Document is not None:
        document_path = st.session_state.Document
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        document_path = os.path.join(script_dir, "example.pdf")
    # Agent 还没创建（首次发问 或 上传新文件后被清空），就创建一个
    if st.session_state.agent is None:
        st.session_state.agent = create_my_agent(document_path)
    # ---------- ⑥ 流式输出到 assistant 气泡 ----------
    with st.chat_message("assistant"):
        full_answer = [""]

        def stream_answer():
            for msg, metadata in st.session_state.agent.stream(
                {"messages": [{"role": "user", "content": question}]},
                config={"configurable": {"thread_id": "ui_session_1"}},
                stream_mode="messages",
            ):
                node = metadata.get("langgraph_node", "")
                if node in ("agent", "model"):
                    content = getattr(msg, "content", "")
                    if content:
                        if isinstance(content, str):
                            full_answer[0] += content
                            yield content
                        elif isinstance(content, list):
                            for part in content:
                                text = part.get("text", "") if isinstance(part, dict) else str(part)
                                if text:
                                    full_answer[0] += text
                                    yield text

        st.write_stream(stream_answer())

    # 回答完成后存入历史
    st.session_state.messages.append({"role": "assistant", "content": full_answer[0]})
